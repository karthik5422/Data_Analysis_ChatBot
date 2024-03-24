import os
import openai
import streamlit as st
from langchain_openai import ChatOpenAI
from langchain.agents.agent_types import AgentType
from streamlit_chat import message
from langchain_experimental.agents.agent_toolkits.pandas.base import create_pandas_dataframe_agent
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

load_dotenv()
openai.api_key = os.getenv('OPENAI_API_KEY')


def initialize_session_state():
    if 'generated' not in st.session_state:
        st.session_state['generated'] = []
    if 'past' not in st.session_state:
        st.session_state['past'] = []
    if 'messages' not in st.session_state:
        st.session_state['messages'] = [{"role": "system", "content": "You are a helpful assistant."}]

def chatbox_functionality(dataframes, uploaded_file_names):
    initialize_session_state()

    llm = ChatOpenAI(temperature=0, model="gpt-3.5-turbo-0613")

    if dataframes:
        selected_file_name = st.selectbox("Select file for chatting:", options=uploaded_file_names)

        if selected_file_name:
            idx = uploaded_file_names.index(selected_file_name)
            pandas_agent = create_pandas_dataframe_agent(llm, dataframes[idx], verbose=True,
                                                          agent_type=AgentType.OPENAI_FUNCTIONS,
                                                          agent_executor_kwargs={"handle_parsing_errors": True})

            st.markdown("<h3 style='text-align: center;'>CHAT WITH BOT</h3>", unsafe_allow_html=True)

            def generate_response(input_text, dataframes):
                response = ""
                response += "**Response for Selected File:**\n"
                response += pandas_agent.run(input_text) + "\n"

                return response

            response_container = st.container()
            download_button_container = st.container()
            input_container = st.container()

            with input_container:
                with st.form(key='my_form', clear_on_submit=True):
                    user_input = st.text_area("**You:**", key='input', height=st.session_state.get('input_height', 50))
                    st.session_state['input_height'] = len(user_input.split('\n')) * 20

                    submit_button = st.form_submit_button(label='Send')

                if submit_button and user_input:
                    try:
                        response = generate_response(user_input, dataframes)
                        st.session_state['past'].append(user_input)
                        st.session_state['generated'].append(response)
                    except Exception as e:
                        st.error("An error occurred: {}".format(e))

            if st.session_state.get('generated'):
                with response_container:
                    st.markdown("---")
                    for i in range(len(st.session_state['generated'])):
                        message(st.session_state["past"][i], is_user=True, key=str(i) + '_user')
                        message(st.session_state["generated"][i], key=str(i))

                with download_button_container:
                    download_button_clicked = False
                    download_button = st.button("Download Chat History")

                    if download_button and not download_button_clicked:
                        download_button_clicked = True
                        chat_history_file_path = "chat_history.docx"
                        document = Document()
                        document.add_heading('Chat History', level=1)
                        document.add_paragraph()

                        for user_input, generated_response in zip(st.session_state["past"], st.session_state["generated"]):
                            user_paragraph = document.add_paragraph()
                            user_paragraph.add_run("You: ").bold = True
                            user_paragraph.add_run(user_input).font.size = Pt(12)

                            bot_paragraph = document.add_paragraph()
                            bot_paragraph.add_run("Bot: ").bold = True
                            bot_paragraph.add_run(generated_response).font.size = Pt(12)

                            document.add_paragraph()

                        document.save(chat_history_file_path)

                        st.markdown(f"[Download Chat History](sandbox:/path/{chat_history_file_path})", unsafe_allow_html=True)
